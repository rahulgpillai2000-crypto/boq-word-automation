#!/usr/bin/env python3
"""Generate a business case Word document from Excel and a Word template.

This script supports two Excel layouts:
1) Tabular layout: one row per project with column headers.
2) Key-value layout: labels and values anywhere across one/multiple sheets.
"""

from __future__ import annotations

import math
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable

import pandas as pd
from docx import Document

EXCEL_PATH = Path("data/master_projects.xlsx")
TEMPLATE_PATH = Path("templates/business_case_template.docx")
OUTPUT_PATH = Path("output/output.docx")

REQUIRED_FIELDS = {
    "Contract_Number",
    "Work_Order_Number",
    "Project_Name",
    "Date",
    "Start_Date",
    "End_Date",
    "Total_Value",
}

FIELD_ALIASES = {
    "Contract_Number": ["contract number", "contract_no", "contract"],
    "Work_Order_Number": ["work order number", "work order", "wo number", "wo"],
    "Project_Name": ["project name", "project"],
    "Date": ["date", "document date"],
    "Start_Date": ["start date", "commencement date", "from date"],
    "End_Date": ["end date", "completion date", "to date"],
    "Total_Value": ["total value", "contract value", "total", "value"],
}

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")


class BusinessCaseError(RuntimeError):
    """Raised for business case generation errors."""


def normalize_key(text: object) -> str:
    if text is None:
        return ""
    cleaned = re.sub(r"[^a-z0-9]+", " ", str(text).strip().lower())
    return re.sub(r"\s+", " ", cleaned).strip()


def _to_datetime(value: object, column_name: str) -> datetime:
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        raise BusinessCaseError(f"Invalid date in field '{column_name}': {value!r}")
    return parsed.to_pydatetime()


def _format_date(value: object) -> str:
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return ""
    return parsed.strftime("%Y-%m-%d")


def _is_blank(value: object) -> bool:
    if value is None:
        return True
    if isinstance(value, float) and pd.isna(value):
        return True
    return str(value).strip() == ""


def candidate_keys_for_field(field_name: str) -> Iterable[str]:
    yield normalize_key(field_name)
    for alias in FIELD_ALIASES.get(field_name, []):
        yield normalize_key(alias)


def extract_from_tabular_layout(xls: pd.ExcelFile) -> Dict[str, object] | None:
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        if df.empty:
            continue

        renamed = {str(col): str(col).strip() for col in df.columns}
        df = df.rename(columns=renamed)
        if REQUIRED_FIELDS.issubset(set(df.columns)):
            row = df.iloc[0]
            return {field: row[field] for field in REQUIRED_FIELDS}
    return None


def extract_from_key_value_layout(xls: pd.ExcelFile) -> Dict[str, object] | None:
    key_to_value: Dict[str, object] = {}

    for sheet in xls.sheet_names:
        raw = pd.read_excel(xls, sheet_name=sheet, header=None)
        if raw.empty:
            continue

        rows, cols = raw.shape
        for r in range(rows):
            for c in range(cols):
                key_cell = raw.iat[r, c]
                key = normalize_key(key_cell)
                if not key:
                    continue

                # Candidate value: right cell first, then below cell.
                value_right = raw.iat[r, c + 1] if c + 1 < cols else None
                value_down = raw.iat[r + 1, c] if r + 1 < rows else None

                value = None
                if not _is_blank(value_right):
                    value = value_right
                elif not _is_blank(value_down):
                    value = value_down

                if value is not None and key not in key_to_value:
                    key_to_value[key] = value

    if not key_to_value:
        return None

    resolved: Dict[str, object] = {}
    for field in REQUIRED_FIELDS:
        for candidate in candidate_keys_for_field(field):
            if candidate in key_to_value:
                resolved[field] = key_to_value[candidate]
                break

    if REQUIRED_FIELDS.issubset(set(resolved.keys())):
        return resolved
    return None


def load_project_data(excel_path: Path) -> Dict[str, object]:
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    xls = pd.ExcelFile(excel_path)

    extracted = extract_from_tabular_layout(xls)
    if extracted is None:
        extracted = extract_from_key_value_layout(xls)

    if extracted is None:
        raise BusinessCaseError(
            "Could not locate all required fields in Excel. Supported layouts are: "
            "(1) tabular columns, or (2) key-value labels across sheets."
        )

    missing = REQUIRED_FIELDS - set(extracted.keys())
    if missing:
        raise BusinessCaseError(f"Missing required fields: {', '.join(sorted(missing))}")

    return extracted


def compute_cashflow(data: Dict[str, object]) -> Dict[str, float | int]:
    start_date = _to_datetime(data["Start_Date"], "Start_Date")
    end_date = _to_datetime(data["End_Date"], "End_Date")
    if end_date < start_date:
        raise BusinessCaseError("End_Date cannot be earlier than Start_Date.")

    duration_days = (end_date - start_date).days + 1
    duration_months = max(1, math.ceil(duration_days / 30))

    try:
        total_value = float(str(data["Total_Value"]).replace(",", ""))
    except ValueError as exc:
        raise BusinessCaseError(f"Invalid Total_Value: {data['Total_Value']!r}") from exc

    monthly_fee = total_value / duration_months
    return {
        "duration_days": duration_days,
        "duration_months": duration_months,
        "total_value": total_value,
        "monthly_fee": monthly_fee,
    }


def replace_simple_placeholders(document: Document, replacements: Dict[str, str]) -> int:
    replaced_count = 0

    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                replaced_count += 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            replaced_count += 1

    return replaced_count


def insert_cashflow_table(document: Document, duration_months: int, monthly_fee: float) -> None:
    marker = "{{CASHFLOW_TABLE}}"

    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            paragraph.text = paragraph.text.replace(marker, "")
            table = document.add_table(rows=2, cols=duration_months)
            table.style = "Table Grid"

            for month in range(duration_months):
                table.cell(0, month).text = f"Month {month + 1}"
                table.cell(1, month).text = f"{monthly_fee:,.2f}"

            paragraph._p.addnext(table._tbl)
            return

    raise BusinessCaseError("Placeholder {{CASHFLOW_TABLE}} not found in template.")


def validate_template_has_placeholders(document: Document) -> None:
    found = set()
    for paragraph in document.paragraphs:
        found.update(PLACEHOLDER_PATTERN.findall(paragraph.text or ""))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    found.update(PLACEHOLDER_PATTERN.findall(paragraph.text or ""))

    expected = {"Contract_Number", "Work_Order_Number", "Project_Name", "Date", "Total_Value", "CASHFLOW_TABLE"}
    if not found.intersection(expected):
        raise BusinessCaseError(
            "No placeholders were found in the template. Add placeholders like "
            "{{Contract_Number}} and {{CASHFLOW_TABLE}} in Word where you want values/table inserted."
        )


def main() -> None:
    data = load_project_data(EXCEL_PATH)
    cashflow = compute_cashflow(data)

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Word template not found: {TEMPLATE_PATH}")

    document = Document(TEMPLATE_PATH)
    validate_template_has_placeholders(document)

    replacements = {
        "Contract_Number": str(data["Contract_Number"]),
        "Work_Order_Number": str(data["Work_Order_Number"]),
        "Project_Name": str(data["Project_Name"]),
        "Date": _format_date(data["Date"]),
        "Total_Value": f"{cashflow['total_value']:,.2f}",
    }

    replace_simple_placeholders(document, replacements)
    insert_cashflow_table(document, int(cashflow["duration_months"]), float(cashflow["monthly_fee"]))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)

    print(f"Generated Word document: {OUTPUT_PATH}")
    print(
        f"Duration days={cashflow['duration_days']}, months={cashflow['duration_months']}, "
        f"monthly fee={cashflow['monthly_fee']:,.2f}"
    )


if __name__ == "__main__":
    main()
