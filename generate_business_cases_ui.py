def generate_document(excel_file, word_template):

    import pandas as pd
    from docx import Document
    from math import ceil

    # =========================
    # LOAD EXCEL
    # =========================
    df = pd.read_excel(excel_file)
    data = df.iloc[0]

    # =========================
    # EXTRACT VALUES
    # =========================
    contract_number = str(data["Contract Number"])
    work_order = str(data["Work Order Number"])
    project_name = str(data["Project Name"])

    # Business Case Date
    date_value = pd.to_datetime(data["Date"])
    formatted_date = date_value.strftime("%d %B %Y")

    # Project Dates
    start_date = pd.to_datetime(data["Commencement Date"])
    end_date = pd.to_datetime(data["End Date"])

    total_value = float(data["Total Value"])

    # =========================
    # CASHFLOW (DAY-BASED SMART)
    # =========================
    duration_days = (end_date - start_date).days

    if duration_days <= 0:
        duration_days = 1

    # Split duration
    full_months = duration_days // 30
    remaining_days = duration_days % 30

    months = full_months + (1 if remaining_days > 0 else 0)

    values = []

    # Value per day
    value_per_day = total_value / duration_days

    # Full months (equal)
    full_month_value = round(value_per_day * 30, 2)

    for _ in range(full_months):
        values.append(full_month_value)

    # Last partial month (smaller)
    if remaining_days > 0:
        last_value = round(total_value - sum(values), 2)
        values.append(last_value)

    # Final rounding correction
    if len(values) > 0:
        values = [round(v, 2) for v in values[:-1]] + [
            round(total_value - sum(values[:-1]), 2)
        ]

    months = len(values)

    # =========================
    # LOAD WORD TEMPLATE
    # =========================
    doc = Document(word_template)

    # =========================
    # BASE PLACEHOLDERS
    # =========================
    mapping = {
        "{{Contract_Number}}": contract_number,
        "{{Work_Order_Number}}": work_order,
        "{{Project_Name}}": project_name,
        "{{Total_Value}}": f"AED {total_value:,.2f}",
        "{{Date}}": formatted_date
    }

    # =========================
    # MONTH PLACEHOLDERS
    # =========================
    for i in range(months):
        mapping[f"{{{{M{i+1}}}}}"] = f"{values[i]:,.2f}"

    for i in range(months, 12):
        mapping[f"{{{{M{i+1}}}}}"] = ""

    # =========================
    # REPLACEMENT FUNCTIONS
    # =========================
    def replace_text_in_paragraph(paragraph):
        for key, val in mapping.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)

    def replace_text_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

    # =========================
    # APPLY REPLACEMENTS
    # =========================
    for p in doc.paragraphs:
        replace_text_in_paragraph(p)

    for table in doc.tables:
        replace_text_in_table(table)

    # =========================
    # SAVE OUTPUT
    # =========================
    output_path = "output/output.docx"
    doc.save(output_path)

    print("Duration Days:", duration_days)
    print("Months:", months)
    print("Values:", values)
    print("Total:", sum(values))

    print("✅ SUCCESS – Business Case Generated")

    return output_path